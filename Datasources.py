import uuid
import os
from pathlib import Path
from typing import List, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import pdfplumber
import docx
import pyodbc
import pandas as pd
import chromadb
import pyodbc
import pandas as pd
from chromadb.utils import embedding_functions
from pathlib import Path
from Functions.classLibrary import Document, Prompt
import os
import mammoth
import importlib
try:
    _sa = importlib.import_module("sqlalchemy")
    create_engine = _sa.create_engine
    text = _sa.text
except ModuleNotFoundError:  # SQLAlchemy not installed
    create_engine = None
    def text(x):
        return x
from urllib.parse import quote_plus

class ELabJobsDB:
    """Handles SQL Server database connections."""

    def __init__(self, server="hchwusrv2062", database="ELabJobs", username="ELabReader",
                 password="sfg7as*TClkc89", driver="{ODBC Driver 17 for SQL Server}"):
        self.config = {
            "server": server,
            "database": database,
            "username": username,
            "password": password,
            "driver": driver
        }

    def get_connection(self):
        """Creates and returns a new raw pyodbc connection (legacy use)."""
        conn_str = (
            f"DRIVER={self.config['driver']};"
            f"SERVER={self.config['server']};"
            f"DATABASE={self.config['database']};"
            f"UID={self.config['username']};"
            f"PWD={self.config['password']}"
        )
        return pyodbc.connect(conn_str)

    def get_sqlalchemy_engine(self):
        """Creates a SQLAlchemy engine for pandas.read_sql to avoid warnings.
        Returns an engine or raises ImportError if SQLAlchemy is not available.
        """
        if create_engine is None:
            raise ImportError("SQLAlchemy is not installed")
        # Build ODBC connection string and URL-encode it for SQLAlchemy
        odbc_parts = [
            f"DRIVER={self.config['driver']}",
            f"SERVER={self.config['server']}",
            f"DATABASE={self.config['database']}",
            f"UID={self.config['username']}",
            f"PWD={self.config['password']}",
            "TrustServerCertificate=yes",
        ]
        odbc_conn_str = ";".join(odbc_parts)
        connect_url = f"mssql+pyodbc:///?odbc_connect={quote_plus(odbc_conn_str)}"
        return create_engine(connect_url)

    def runQuery(self, prompt: str) -> pd.DataFrame:
        """
        Runs a SQL query against the database using SQLAlchemy engine.
        Returns a pandas DataFrame with the results.
        """
        try:
            engine = self.get_sqlalchemy_engine()
            with engine.begin() as conn:
                df = pd.read_sql(text(prompt), conn)
                return df
        except Exception as e:
            # Fallback to legacy connection if SQLAlchemy path fails
            try:
                import warnings
                with self.get_connection() as legacy_conn:
                    with warnings.catch_warnings():
                        # Suppress the pandas UserWarning about DBAPI2; we prefer SQLAlchemy but continue gracefully
                        warnings.simplefilter("ignore", category=UserWarning)
                        df = pd.read_sql(prompt, legacy_conn)
                    return df
            except Exception:
                pass
            print(f"❌ SQL Error in runQuery: {e}")
            return pd.DataFrame([{"error": str(e)}])

class LocalFolderSource:
    """Source-Implementierung für lokale Ordner mit Dokumenten (PDF, DOCX, TXT)."""

    def __init__(self, folder_path: str):
        self.folder_path = Path(folder_path)
        if not self.folder_path.exists():
            raise FileNotFoundError(f"Ordner {folder_path} existiert nicht.")
        self.connected = False

    def connect(self):
        """Prüft, ob der Ordner zugänglich ist."""
        if not os.access(self.folder_path, os.R_OK):
            raise PermissionError(f"Keine Leserechte für {self.folder_path}")
        self.connected = True
        print(f"📁 Verbunden mit lokalem Ordner: {self.folder_path}")

    def disconnect(self):
        """Schließt die (symbolische) Verbindung."""
        self.connected = False
        print("🔌 Lokale Ordnerverbindung getrennt")

    def _extract_text_from_file(
            self,
            file_path: Path,
            mode: str = "full",
            lines_per_chunk: int = 80,
            chars_per_chunk: int = 1000
    ) -> Tuple[str, List[str]]:

        mode_map = {
            "seiten": "pages",
            "zeilen": "file",
            "ganzes_dokument": "full",
            "zeichen": "chars"
        }
        mode = mode_map.get(mode.lower(), "full")
        ext = file_path.suffix.lower()

        try:
            # ------------------------------
            # 🔹 PDF
            # ------------------------------
            if ext == ".pdf":
                with pdfplumber.open(file_path) as pdf:
                    pages = [(p.extract_text() or "").strip() for p in pdf.pages]

                text = "\n".join(pages)

                if mode == "pages":
                    chunks = pages
                elif mode == "file":
                    chunks = self._chunk_text_by_lines(text, lines_per_chunk)
                elif mode == "chars":
                    chunks = self._chunk_text_by_chars(text, chars_per_chunk)
                else:  # full
                    chunks = [text]

                return text.strip(), chunks

            # ------------------------------
            # 🔹 DOCX
            # ------------------------------
            elif ext == ".docx":
                paragraphs = [p.text for p in docx.Document(file_path).paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
                print(text)

            # ------------------------------
            # 🔹 Textbasierte Formate
            # ------------------------------
            elif ext in {".txt", ".csv", ".log"}:
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    text = f.read()

            # ------------------------------
            # 🔹 Nicht unterstützt
            # ------------------------------
            else:
                text = f"[Nicht unterstützter Dateityp: {ext}]"
                return text, [text]

        except Exception as e:
            text = f"[Fehler beim Lesen: {e}]"
            return text, [text]

        # ------------------------------
        # 🔹 Chunk-Verarbeitung für Nicht-PDFs
        # ------------------------------
        if mode == "pages":
            # Für DOCX/TXT etc. als "Seiten" durch Absätze definieren
            chunks = text.split("\n\n")  # doppelte neue Zeilen als Absatztrenner
        elif mode == "file":
            chunks = self._chunk_text_by_lines(text, lines_per_chunk)
        elif mode == "chars":
            chunks = self._chunk_text_by_chars(text, chars_per_chunk)
        else:  # full
            chunks = [text]

        return text.strip(), chunks

        # ------------------------------
        # 🔹 Chunking Helper
        # ------------------------------

    def _chunk_text_by_lines(self, text: str, lines_per_chunk: int = 80) -> List[str]:
        """Teilt Text in Chunks von N Zeilen auf."""
        lines = text.splitlines()
        chunks = []
        for i in range(0, len(lines), lines_per_chunk):
            chunk = "\n".join(lines[i:i + lines_per_chunk]).strip()
            if chunk:
                chunks.append(chunk)
        return chunks or [""]

    def _chunk_text_by_chars(self, text: str, chars_per_chunk: int = 1000) -> List[str]:
        """Teilt Text in Chunks von N Zeichen auf."""
        chunks = []
        for i in range(0, len(text), chars_per_chunk):
            chunk = text[i:i + chars_per_chunk].strip()
            if chunk:
                chunks.append(chunk)
        return chunks or [""]

    # ------------------------------
    # 🔹 Dokumente laden
    # ------------------------------
    def get_documents(self, topN: int = 100, extensions: Optional[List[str]] = None) -> List[Document]:
        """
        Lädt bis zu `topN` Dokumente aus dem lokalen Ordner, extrahiert Text und erstellt Chunks.
        Unterstützte Typen: .pdf, .docx, .txt (standardmäßig)
        """
        if not self.connected:
            raise ConnectionError("Nicht verbunden. Bitte zuerst `connect()` aufrufen.")

        extensions = extensions or [".pdf", ".docx", ".txt", ".doc"]
        all_files = [f for f in self.folder_path.rglob("*") if f.suffix.lower() in extensions]
        documents = []
        for i, file_path in enumerate(all_files[:topN]):
            print(file_path)
            content, chunks = self._extract_text_from_file(file_path, mode="ganzes_dokument")
            created = datetime.fromtimestamp(file_path.stat().st_ctime).isoformat()
            modified = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            size = file_path.stat().st_size

            documents.append(Document(
                id=str(uuid.uuid4()),
                title=file_path.stem,
                location=str(file_path),
                tags=[file_path.suffix.lower().lstrip(".")],
                created=created,
                content=content,
                chunks=chunks,
            ))

        return documents

    # ------------------------------
    # 🔹 Einzelnes Dokument lesen
    # ------------------------------
    def get_document_content(self, doc_path: str) -> tuple[str, dict]:
        """Lädt den Inhalt und einfache Metadaten eines Dokuments."""
        file_path = Path(doc_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Datei {doc_path} nicht gefunden.")

        content, chunks = self._extract_text_from_file(file_path)
        created = datetime.fromtimestamp(file_path.stat().st_ctime).isoformat()
        modified = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        size = file_path.stat().st_size

        metadata = {
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower(),
            "size_bytes": size,
            "created": created,
            "modified": modified,
            "pages_or_chunks": len(chunks)
        }
        return content, metadata


class ChromaDB:
    """
    Wrapper um Chroma mit OpenAI Embeddings.
    Arbeitet direkt mit Document-Objekten, die Chunks und Metadaten enthalten.
    """

    def __init__(self, collection: str, persist_directory: str = "./chroma_db") -> None:
        self.client = chromadb.PersistentClient(path=persist_directory)

        self.embedding_function = embedding_functions.OpenAIEmbeddingFunction(
            api_key=os.getenv("OPENAI_API_KEY"),
            model_name="text-embedding-3-large"
        )

        self.collection = self.client.get_or_create_collection(
            name=collection,
            embedding_function=self.embedding_function
        )

    # --------------------------------------------------
    # Verwaltung
    # --------------------------------------------------
    def list_collections(self):
        return self.client.list_collections()

    def remove_file(self, doc_id: str):
        self.collection.delete(ids=[doc_id])

    def drop_db(self):
        try:
            self.client.delete_collection(self.collection.name)
        except(Exception):
            print("No Collection to delete")


    # --------------------------------------------------
    # Dokumente hinzufügen
    # --------------------------------------------------
    def add_document(self, document: Document):
        """
        Fügt ein Document-Objekt mit seinen Chunks in die Chroma-Collection ein.
        Jeder Chunk wird als eigener Eintrag gespeichert.
        Die Metadaten werden aus Document.metadata übernommen.
        """
        if not document.chunks or len(document.chunks) == 0:
            raise ValueError(f"❌ Document '{document.title}' enthält keine Chunks.")

        texts = document.chunks
        ids = [str(uuid.uuid4()) for _ in texts]

        # Tags sicher als String speichern
        tags = None
        if document.tags:
            if isinstance(document.tags, list):
                tags = ",".join(map(str, document.tags))
            else:
                tags = str(document.tags)

        # Basis-Metadaten aus dem Document
        base_meta = {
            "document_id": str(document.id),
            "document_title": document.title,
            "document_location": document.location,
            "created": document.created,
            "tags": tags,
        }

        # === Inhaltliche Metadaten aus MetaData ===
        if document.meta:
            md = document.meta

            # alle Listen-Felder (keywords, authors, recipient, linked_Documents)
            # müssen in Strings konvertiert werden (z. B. JSON)
            def safe_list(val):
                if val is None:
                    return None
                if isinstance(val, list):
                    return ", ".join(map(str, val))
                return str(val)

            base_meta.update({
                "doc_type": md.doc_type,
                "language": md.language,
                "authors": safe_list(md.authors),
                "abstract": md.abstract,
                "keywords": safe_list(md.keywords),
                "sender": md.sender,
                "recipient": safe_list(md.recipient),
                "body": md.body,
                "huba_document": bool(md.huba_document) if md.huba_document is not None else None,
                "linked_documents": safe_list(md.linked_Documents),
                "product_code": md.product_code,
                "art_no": md.art_no,
            })

        # Für jeden Chunk eine Kopie der Metadaten mit Chunk-Index
        metadatas = []
        for i, _ in enumerate(texts):
            chunk_meta = base_meta.copy()
            chunk_meta["chunk_index"] = i
            metadatas.append(chunk_meta)

        try:
            # Embeddings erzeugen
            embeddings = self.embedding_function(texts)

            # In Chroma einfügen
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=texts,
                metadatas=metadatas
            )
        except Exception as e:
            print("Failed to add Chunk")
            print (e)

        print(f"✅ {len(texts)} Chunks von '{document.title}' mit MetaData hinzugefügt.")

    # --------------------------------------------------
    # Laden & Abfragen
    # --------------------------------------------------
    def load_files(self) -> pd.DataFrame:
        data = self.collection.get(include=["documents", "metadatas"])
        documents = data["documents"]
        metadatas = data["metadatas"]
        rows = []
        for i in range(len(documents)):
            row = {"content": documents[i]}
            if metadatas[i] is not None:
                row.update(metadatas[i])
            rows.append(row)

        return pd.DataFrame(rows)

    def semanticRetrieval(self, prompt:Prompt, n_results: int = 2):
        results = self.collection.query(
            query_texts=[prompt.user_prompt],
            n_results=n_results,
            include=["documents"]
        )
        return results


    def keywordsRetrieval(self, keywords: list[str], n_results: int = 2):
        """Filter documents by provided keywords/product codes across multiple metadata fields.
        Returns a pandas DataFrame with at most n_results rows.
        """
        df = self.load_files()
        try:
            import re
            import pandas as _pd  # local alias
        except Exception:
            return df.head(0)

        if df is None or df.empty:
            return df
        if not keywords:
            return df.head(0)

        # Build a case-insensitive regex combining all keywords safely
        pats = [re.escape(str(k)) for k in keywords if str(k).strip()]
        if not pats:
            return df.head(0)
        regex = "|".join(pats)

        # Candidate columns to match against
        candidate_cols = [
            "keywords",
            "product_code",
            "document_title",
            "document_location",
            "title",
            "art_no",
            "content",
        ]
        mask = _pd.Series(False, index=df.index)
        for col in candidate_cols:
            if col in df.columns:
                try:
                    mask = mask | df[col].astype(str).str.contains(regex, case=False, na=False)
                except Exception:
                    continue
        result = df[mask]
        if result.empty:
            return result
        # Return top n_results (no specific scoring available here)
        return result.head(max(1, int(n_results)))